import docx
import os
import pandas as pd

# 实验报告所在路径
base = "/run/media/znw/Ventoy/刻盘/数据可视化-2022级计算机科学与技术（专升本）5班-卓能文/9.实验资料/1.平时实验资料/实验1- ECharts数据可视化（一）"

data = pd.read_excel("data.xlsx")
for i, datarow in data.iterrows():
    path = str(datarow.学号) + "-" + datarow.姓名
    filename = base + "/" + path + "/" + path + ".docx"
    if os.path.exists(filename):
        try:
            doc = docx.Document(filename)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.startswith("成绩评定"):  # 检查单元格内容
                            cell.text = ""  # 清空单元格内容
                            cell.paragraphs[0].add_run("成绩评定：").bold = True
                            cell.paragraphs[0].add_run("\n" + str(datarow.成绩))
                            doc.save(filename)
        except Exception as e:
            print(path + " 文件读取失败。原因： " + str(e))
        finally:
            continue
    else:
        print(path + " 不存在")
